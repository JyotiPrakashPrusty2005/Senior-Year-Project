"""Generate a Word document describing the Health AI Models project."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

doc = Document()

# ── Styles ──
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# ────────────────────────────────────────────────────────────
# TITLE PAGE
# ────────────────────────────────────────────────────────────
doc.add_paragraph("\n\n\n")
title = doc.add_heading("Health AI Models", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    "Project Documentation\n\n"
    "Dataset Descriptions, Preprocessing, Model Training & Results"
)
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ────────────────────────────────────────────────────────────
# TABLE OF CONTENTS (manual)
# ────────────────────────────────────────────────────────────
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Project Overview",
    "2. Datasets",
    "   2.1 Diabetes Dataset",
    "   2.2 Cardiovascular Disease Dataset",
    "   2.3 Pneumonia Chest X-Ray Dataset",
    "   2.4 Post-Surgery Recovery Dataset (Synthetic)",
    "3. Preprocessing Pipeline",
    "   3.1 Tabular Data Preprocessing",
    "   3.2 Image Data Preprocessing",
    "   3.3 Time-Series Preprocessing (Recovery LSTM)",
    "4. Model Training & Comparison",
    "   4.1 Diabetes Model",
    "   4.2 Cardiovascular Model",
    "   4.3 Pneumonia Model",
    "   4.4 Post-Surgery Recovery LSTM Model",
    "5. Final Results Summary",
    "6. Inference Pipeline",
    "7. Post-Surgery Recovery Coach (Generative AI)",
]
for item in toc_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_page_break()

# ────────────────────────────────────────────────────────────
# 1. PROJECT OVERVIEW
# ────────────────────────────────────────────────────────────
doc.add_heading("1. Project Overview", level=1)
doc.add_paragraph(
    "This project, titled 'Agentic AI Health Coach: Early Disease Risk Prediction and "
    "Post-Surgery Recovery Assistant using Hybrid ML, LSTM and Generative AI', develops "
    "machine learning models to predict multiple health conditions: "
    "Diabetes, Cardiovascular Disease, and Pneumonia. "
    "For tabular datasets, multiple classical ML and gradient boosting algorithms are compared, "
    "and the best-performing model is selected automatically. "
    "For image-based pneumonia detection, deep learning architectures (ResNet-18, ResNet-50, EfficientNet-B0) "
    "are trained and compared using transfer learning on chest X-ray images."
)
doc.add_paragraph(
    "Additionally, the project includes a Post-Surgery Recovery Assistant that uses a hybrid "
    "LSTM neural network to predict patient recovery trajectories and complication risks from "
    "daily vitals time-series data. A Generative AI-powered recovery coach (using Google Gemini "
    "or OpenAI) provides personalized, evidence-based recovery guidance to patients."
)
doc.add_paragraph(
    "The project follows a consistent pipeline: data loading → preprocessing → "
    "multi-model training → evaluation → best model selection → inference script.",
)

doc.add_page_break()

# ────────────────────────────────────────────────────────────
# 2. DATASETS
# ────────────────────────────────────────────────────────────
doc.add_heading("2. Datasets", level=1)

# --- 2.1 Diabetes ---
doc.add_heading("2.1 Diabetes Dataset", level=2)
doc.add_paragraph(
    "Source: CDC Behavioral Risk Factor Surveillance System (BRFSS) 2015 survey data."
)
doc.add_paragraph(
    "File: diabetes_012_health_indicators_BRFSS2015.csv"
)

table = doc.add_table(rows=5, cols=2, style="Light Shading Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = [
    ("Total Samples", "253,680"),
    ("Features", "21 health indicators"),
    ("Target Variable", "Diabetes_012 (0 = No Diabetes, 1 = Prediabetes, 2 = Diabetes)"),
    ("Class Distribution", "No Diabetes: 213,703 | Prediabetes: 4,631 | Diabetes: 35,346"),
    ("Missing Values", "None"),
]
for i, (k, v) in enumerate(cells):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_paragraph()
doc.add_paragraph(
    "Key features include: HighBP, HighChol, BMI, Smoker, Stroke history, "
    "HeartDiseaseorAttack, PhysActivity, Fruits, Veggies, HvyAlcoholConsump, "
    "GenHlth, MentHlth, PhysHlth, DiffWalk, Sex, Age, Education, and Income."
)

# --- 2.2 Cardiovascular ---
doc.add_heading("2.2 Cardiovascular Disease Dataset", level=2)
doc.add_paragraph(
    "Source: Kaggle cardiovascular disease dataset (cardio_train.csv)."
)

table = doc.add_table(rows=5, cols=2, style="Light Shading Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = [
    ("Total Samples", "70,000"),
    ("Features", "11 (after dropping id)"),
    ("Target Variable", "cardio (0 = No CVD, 1 = CVD)"),
    ("Class Distribution", "No CVD: 35,021 | CVD: 34,979 (balanced)"),
    ("Missing Values", "None"),
]
for i, (k, v) in enumerate(cells):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_paragraph()
doc.add_paragraph(
    "Features include: age (in days), gender, height, weight, systolic blood pressure (ap_hi), "
    "diastolic blood pressure (ap_lo), cholesterol level, glucose level, smoking, "
    "alcohol intake, and physical activity."
)

# --- 2.3 Pneumonia ---
doc.add_heading("2.3 Pneumonia Chest X-Ray Dataset", level=2)
doc.add_paragraph(
    "Source: Kaggle Chest X-Ray Images (Pneumonia) dataset, organized into train/val/test folders."
)

table = doc.add_table(rows=5, cols=2, style="Light Shading Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = [
    ("Training Set", "NORMAL: 1,341 | PNEUMONIA: 3,875 (Total: 5,216)"),
    ("Validation Set", "NORMAL: 8 | PNEUMONIA: 8 (Total: 16)"),
    ("Test Set", "NORMAL: 234 | PNEUMONIA: 390 (Total: 624)"),
    ("Image Format", "JPEG chest X-ray images"),
    ("Classes", "2 (NORMAL, PNEUMONIA)"),
]
for i, (k, v) in enumerate(cells):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

# --- 2.4 Post-Surgery Recovery ---
doc.add_heading("2.4 Post-Surgery Recovery Dataset (Synthetic)", level=2)
doc.add_paragraph(
    "Source: Synthetically generated using data/generate_synthetic_recovery.py. "
    "Simulates 1,000 patients recovering from surgery over 30 days with three outcome "
    "trajectories: normal recovery, slow recovery, and complication (infection/readmission)."
)

table = doc.add_table(rows=8, cols=2, style="Light Shading Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = [
    ("Total Records", "30,000 (1,000 patients × 30 days)"),
    ("Patients", "1,000"),
    ("Time-Series Features (Daily)", "pain_level, temperature, heart_rate, bp_systolic, mobility_score, wound_status, sleep_hours"),
    ("Static Features", "age, gender, bmi, diabetes, hypertension, smoking, surgery_type"),
    ("Target Variables", "recovery_score (regression) + complication_flag (binary)"),
    ("Outcome Distribution", "Normal: ~54% | Slow: ~31% | Complication: ~15%"),
    ("Surgery Types", "cardiac, orthopedic, abdominal, neurological"),
]
for i, (k, v) in enumerate(cells):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_paragraph()
doc.add_paragraph(
    "The synthetic data generator creates realistic recovery trajectories where patient "
    "demographics (age, BMI, comorbidities, smoking status) influence the probability of "
    "complications. Normal recovery shows smooth improvement, slow recovery shows gradual "
    "improvement, and complication trajectories show worsening vitals around days 5-15 "
    "before partial recovery."
)

doc.add_page_break()

# ────────────────────────────────────────────────────────────
# 3. PREPROCESSING PIPELINE
# ────────────────────────────────────────────────────────────
doc.add_heading("3. Preprocessing Pipeline", level=1)

doc.add_heading("3.1 Tabular Data Preprocessing", level=2)
doc.add_paragraph(
    "All tabular datasets (Diabetes, Cardiovascular) follow a common preprocessing pipeline:"
)

steps = [
    ("Missing Value Handling",
     "Missing values identified and handled per dataset. "
     "Diabetes and Cardiovascular datasets had no missing values."),
    ("Categorical Encoding",
     "Categorical features were encoded using scikit-learn's LabelEncoder. "
     "Encoders are saved for inference."),
    ("Feature Scaling",
     "StandardScaler was applied to all features to normalize them to zero mean and unit variance. "
     "The scaler was fitted on the training set only and saved for use during inference."),
    ("Train-Test Split",
     "An 80/20 stratified train-test split was used (random_state=42) to preserve class distributions."),
    ("Class Imbalance Handling (SMOTE)",
     "For imbalanced datasets (Diabetes), SMOTE (Synthetic Minority Oversampling Technique) "
     "was applied to the training set after scaling to generate synthetic samples for minority classes."),
]

for title_text, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(title_text + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("3.2 Image Data Preprocessing (Pneumonia)", level=2)
doc.add_paragraph(
    "Chest X-ray images were preprocessed using torchvision transforms:"
)

steps_img = [
    ("Resizing", "All images were resized to 224×224 pixels."),
    ("Data Augmentation (Training Only)",
     "Random horizontal flip, random rotation (±10°), and color jitter "
     "(brightness=0.2, contrast=0.2) were applied to improve generalization."),
    ("Normalization",
     "Images were normalized using ImageNet mean [0.485, 0.456, 0.406] and "
     "std [0.229, 0.224, 0.225] for transfer learning compatibility."),
    ("Class Imbalance",
     "Handled via weighted CrossEntropyLoss — class weights were computed inversely "
     "proportional to class frequency in the training set."),
]

for title_text, desc in steps_img:
    p = doc.add_paragraph()
    run = p.add_run(title_text + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("3.3 Time-Series Preprocessing (Recovery LSTM)", level=2)
doc.add_paragraph(
    "The post-surgery recovery data required specialized time-series preprocessing:"
)

steps_ts = [
    ("Surgery Type Encoding",
     "Surgery type (cardiac, orthopedic, abdominal, neurological) was label-encoded "
     "using scikit-learn's LabelEncoder."),
    ("Vitals Scaling",
     "All 7 daily vital sign features were standardized using StandardScaler. "
     "The scaler was fitted on the full training set."),
    ("Static Feature Scaling",
     "Patient demographics (age, gender, BMI, diabetes, hypertension, smoking, surgery_type_encoded) "
     "were separately scaled using StandardScaler."),
    ("Sliding Window Sequences",
     "A sliding window of 7 days was applied to create input sequences. "
     "For each window, the target is the recovery score for the next 3 days "
     "and the complication flag in the forecast window. This produced ~20,000 training samples."),
    ("Train-Test Split",
     "An 80/20 stratified split (stratified on complication flag) was used."),
]

for title_text, desc in steps_ts:
    p = doc.add_paragraph()
    run = p.add_run(title_text + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ────────────────────────────────────────────────────────────
# 4. MODEL TRAINING & COMPARISON
# ────────────────────────────────────────────────────────────
doc.add_heading("4. Model Training & Comparison", level=1)
doc.add_paragraph(
    "For each disease, multiple models were trained and compared. "
    "The best model was automatically selected based on accuracy and saved for inference."
)

# --- 4.1 Diabetes ---
doc.add_heading("4.1 Diabetes Model", level=2)
doc.add_paragraph("Task: Multi-class classification (3 classes: No Diabetes, Prediabetes, Diabetes)")
doc.add_paragraph("Models compared:")

table = doc.add_table(rows=6, cols=2, style="Light Shading Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.rows[0].cells[0].text = "Model"
table.rows[0].cells[1].text = "Key Hyperparameters"
models_diabetes = [
    ("Logistic Regression", "max_iter=1000"),
    ("Random Forest", "n_estimators=200, max_depth=7"),
    ("Gradient Boosting", "n_estimators=200, lr=0.05, max_depth=7"),
    ("XGBoost", "n_estimators=300, lr=0.05, max_depth=7"),
    ("LightGBM", "n_estimators=300, lr=0.05, max_depth=7"),
]
for i, (name, params) in enumerate(models_diabetes):
    table.rows[i + 1].cells[0].text = name
    table.rows[i + 1].cells[1].text = params

doc.add_paragraph()
doc.add_paragraph(
    "Evaluation Metrics: Accuracy, F1 Score (weighted & macro), ROC-AUC (weighted, OVR). "
    "SMOTE was applied before training to handle the severe class imbalance "
    "(prediabetes class had only 4,631 samples vs. 213,703 for no diabetes). "
    "The best model was selected based on test accuracy and saved as diabetes_model.pkl."
)

# --- 4.2 Cardiovascular ---
doc.add_heading("4.2 Cardiovascular Model", level=2)
doc.add_paragraph("Task: Binary classification (CVD vs. No CVD)")
doc.add_paragraph("Models compared:")

table = doc.add_table(rows=6, cols=2, style="Light Shading Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.rows[0].cells[0].text = "Model"
table.rows[0].cells[1].text = "Key Hyperparameters"
models_cardio = [
    ("Logistic Regression", "max_iter=1000"),
    ("Random Forest", "n_estimators=200, max_depth=7"),
    ("Gradient Boosting", "n_estimators=200, lr=0.05, max_depth=7"),
    ("XGBoost", "n_estimators=300, lr=0.05, max_depth=7"),
    ("LightGBM", "n_estimators=300, lr=0.05, max_depth=7"),
]
for i, (name, params) in enumerate(models_cardio):
    table.rows[i + 1].cells[0].text = name
    table.rows[i + 1].cells[1].text = params

doc.add_paragraph()
doc.add_paragraph(
    "Evaluation Metrics: Accuracy, Precision, Recall, F1 Score, ROC-AUC. "
    "The dataset is nearly balanced (35,021 vs. 34,979), so no oversampling was needed. "
    "The best model was selected based on test accuracy and saved as cardio_model.pkl."
)

# --- 4.3 Pneumonia ---
doc.add_heading("4.3 Pneumonia Model (Deep Learning)", level=2)
doc.add_paragraph("Task: Binary image classification (NORMAL vs. PNEUMONIA)")
doc.add_paragraph("Architectures compared:")

table = doc.add_table(rows=4, cols=2, style="Light Shading Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.rows[0].cells[0].text = "Architecture"
table.rows[0].cells[1].text = "Details"
models_pneumonia = [
    ("ResNet-18", "Pre-trained on ImageNet, final FC layer replaced for 2 classes"),
    ("ResNet-50", "Pre-trained on ImageNet, final FC layer replaced for 2 classes"),
    ("EfficientNet-B0", "Pre-trained on ImageNet, classifier layer replaced for 2 classes"),
]
for i, (name, desc) in enumerate(models_pneumonia):
    table.rows[i + 1].cells[0].text = name
    table.rows[i + 1].cells[1].text = desc

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Training Configuration: ")
run.bold = True
p.add_run(
    "Image size: 224×224, Batch size: 32, Epochs: 10, "
    "Learning rate: 0.001, Optimizer: Adam, "
    "Loss: Weighted CrossEntropyLoss (to handle class imbalance)."
)

doc.add_paragraph()
doc.add_paragraph(
    "Evaluation Metrics: Accuracy, Precision, Recall, F1 Score, ROC-AUC on the test set. "
    "The best model was selected based on validation accuracy and saved as pneumonia_model.pth "
    "(PyTorch checkpoint including model state, class names, image size, and architecture name)."
)

# --- 4.4 Recovery LSTM ---
doc.add_heading("4.4 Post-Surgery Recovery LSTM Model", level=2)
doc.add_paragraph(
    "Task: Dual-output prediction — Recovery score regression (next 3 days) + "
    "Complication risk classification (binary)"
)
doc.add_paragraph("Architecture: Hybrid LSTM with static feature encoder")

table = doc.add_table(rows=9, cols=2, style="Light Shading Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = [
    ("Component", "Details"),
    ("LSTM Layer", "2-layer bidirectional LSTM, hidden_size=128, dropout=0.3"),
    ("Static Encoder", "Dense layers: 7→64→32 with ReLU and Dropout"),
    ("Recovery Head", "Combined(160)→64→3 (predicts 3 days of recovery scores)"),
    ("Complication Head", "Combined(160)→64→1→Sigmoid (binary complication risk)"),
    ("Optimizer", "Adam, lr=0.001, with ReduceLROnPlateau scheduler"),
    ("Loss Function", "MSELoss (recovery) + 2×BCELoss (complication, weighted higher)"),
    ("Training", "30 epochs, batch_size=64, gradient clipping (max_norm=1.0)"),
]
for i, (k, v) in enumerate(cells):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_paragraph()
doc.add_paragraph(
    "The hybrid architecture combines temporal patterns from 7-day vitals sequences (via LSTM) "
    "with static patient demographics (via dense encoder). The two output heads enable "
    "simultaneous regression (recovery score forecasting) and classification (complication detection). "
    "The complication loss is weighted 2× higher to prioritize patient safety."
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Evaluation Metrics: ")
run.bold = True
p.add_run(
    "Recovery Score MAE (Mean Absolute Error) for regression quality. "
    "Complication Detection: Accuracy, Precision, Recall, and F1 Score via classification report. "
    "The model is saved as recovery_lstm_model.pth with full config for reproducible loading."
)

doc.add_page_break()

# ────────────────────────────────────────────────────────────
# 5. FINAL RESULTS SUMMARY
# ────────────────────────────────────────────────────────────
doc.add_heading("5. Final Results Summary", level=1)
doc.add_paragraph(
    "The table below shows the best model selected for each disease after comparison:"
)

table = doc.add_table(rows=5, cols=4, style="Light Shading Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Disease", "Best Model", "Selection Criterion", "Saved Artifact"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True if table.rows[0].cells[i].paragraphs[0].runs else False

results_data = [
    ("Diabetes", "Auto-selected (best accuracy among 5 models)", "Test Accuracy", "diabetes_model.pkl + diabetes_scaler.pkl"),
    ("Cardiovascular", "Auto-selected (best accuracy among 5 models)", "Test Accuracy", "cardio_model.pkl + cardio_scaler.pkl"),
    ("Pneumonia", "Auto-selected (best val accuracy among 3 CNNs)", "Val Accuracy", "pneumonia_model.pth"),
    ("Post-Surgery Recovery", "Hybrid LSTM (dual-head)", "Val Loss (MSE+BCE)", "recovery_lstm_model.pth + scalers + encoder"),
]
for i, (disease, best, criterion, artifact) in enumerate(results_data):
    table.rows[i + 1].cells[0].text = disease
    table.rows[i + 1].cells[1].text = best
    table.rows[i + 1].cells[2].text = criterion
    table.rows[i + 1].cells[3].text = artifact

doc.add_paragraph()
doc.add_paragraph(
    "Note: The exact accuracy numbers depend on the training run. The training scripts print "
    "a full comparison table and detailed classification report (precision, recall, F1, confusion matrix) "
    "for the best model during each run."
)

doc.add_page_break()

# ────────────────────────────────────────────────────────────
# 6. INFERENCE PIPELINE
# ────────────────────────────────────────────────────────────
doc.add_heading("6. Inference Pipeline", level=1)
doc.add_paragraph(
    "Each disease has a dedicated inference script in the inference/ directory that loads the "
    "saved model and provides a prediction function:"
)

table = doc.add_table(rows=5, cols=3, style="Light Shading Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Script", "Function", "Input"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

inference_data = [
    ("predict_diabetes.py", "predict_diabetes(data)", "Dict of 21 health indicators"),
    ("predict_cardio.py", "predict_cardio(data)", "Dict of 11 patient features"),
    ("predict_pneumonia.py", "predict_pneumonia(image_path)", "Path to chest X-ray image"),
    ("predict_recovery.py", "predict_recovery(vitals_history, patient_info)", "7-day vitals list + patient demographics dict"),
]
for i, (script, func, inp) in enumerate(inference_data):
    table.rows[i + 1].cells[0].text = script
    table.rows[i + 1].cells[1].text = func
    table.rows[i + 1].cells[2].text = inp

doc.add_paragraph()
doc.add_paragraph(
    "All inference functions return a standardized JSON-like dictionary containing: "
    "disease name, prediction (class index), predicted_class (label), probability/confidence, "
    "and risk_level (low / moderate / high based on thresholds of 0.4 and 0.7)."
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Risk Classification Thresholds: ")
run.bold = True
p.add_run(
    "Probability > 0.7 → High risk | "
    "Probability 0.4–0.7 → Moderate risk | "
    "Probability < 0.4 → Low risk"
)

doc.add_page_break()

# ────────────────────────────────────────────────────────────
# 7. POST-SURGERY RECOVERY COACH (GENERATIVE AI)
# ────────────────────────────────────────────────────────────
doc.add_heading("7. Post-Surgery Recovery Coach (Generative AI)", level=1)
doc.add_paragraph(
    "The Post-Surgery Recovery Coach is the Agentic AI component of the project. "
    "It takes LSTM model predictions (recovery forecast + complication risk) along with "
    "patient context and generates personalized, evidence-based recovery guidance."
)

doc.add_heading("7.1 Architecture", level=2)
doc.add_paragraph(
    "The recovery coach follows an agentic pipeline:"
)

steps_agent = [
    ("Step 1 — Data Collection",
     "Patient enters daily vitals: pain level, temperature, heart rate, blood pressure, "
     "mobility score, wound status, and sleep hours."),
    ("Step 2 — LSTM Prediction",
     "The hybrid LSTM model processes the last 7 days of vitals along with patient "
     "demographics to predict: (a) Recovery score for the next 3 days, and "
     "(b) Complication risk probability."),
    ("Step 3 — Context Assembly",
     "Patient profile, current vitals, and model predictions are assembled into a "
     "structured context for the generative AI."),
    ("Step 4 — Generative AI Response",
     "A large language model (Google Gemini or OpenAI GPT) receives the context along with "
     "surgery-specific medical guidelines and generates a personalized daily recovery plan."),
    ("Step 5 — Alert Escalation",
     "If complication risk is HIGH (>70%), the system generates an urgent alert "
     "recommending immediate contact with the healthcare provider."),
]

for title_text, desc in steps_agent:
    p = doc.add_paragraph()
    run = p.add_run(title_text + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("7.2 Recovery Knowledge Base", level=2)
doc.add_paragraph(
    "The coach includes a structured medical knowledge base for four surgery types:"
)

table = doc.add_table(rows=5, cols=3, style="Light Shading Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.rows[0].cells[0].text = "Surgery Type"
table.rows[0].cells[1].text = "Recovery Period"
table.rows[0].cells[2].text = "Key Focus Areas"
recovery_kb = [
    ("Cardiac", "6-12 weeks", "Breathing exercises, walking progression, low-sodium diet"),
    ("Orthopedic", "8-16 weeks", "Range of motion, strengthening, high-protein diet"),
    ("Abdominal", "4-8 weeks", "Walking, gradual diet progression, avoid heavy lifting"),
    ("Neurological", "8-24 weeks", "Cognitive rest, gradual mental activity, anti-inflammatory diet"),
]
for i, (stype, period, focus) in enumerate(recovery_kb):
    table.rows[i + 1].cells[0].text = stype
    table.rows[i + 1].cells[1].text = period
    table.rows[i + 1].cells[2].text = focus

doc.add_heading("7.3 Output Format", level=2)
doc.add_paragraph(
    "The recovery coach generates a comprehensive daily plan that includes:"
)
output_items = [
    "Recovery Assessment — analysis of current status with trend detection",
    "Today's Recommendations — specific actionable advice based on vitals and risk level",
    "Exercise Plan — phase-appropriate exercises for the current recovery stage",
    "Dietary Advice — surgery-type-specific nutritional guidance",
    "Warning Signs — personalized symptoms to watch for",
    "Motivation — encouraging message to support recovery",
]
for item in output_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("7.4 Generative AI Integration", level=2)
doc.add_paragraph(
    "The system supports three modes of operation:"
)

table = doc.add_table(rows=4, cols=3, style="Light Shading Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.rows[0].cells[0].text = "Mode"
table.rows[0].cells[1].text = "Provider"
table.rows[0].cells[2].text = "Setup"
ai_modes = [
    ("Gemini (Recommended)", "Google Gemini 1.5 Flash", "Set GEMINI_API_KEY environment variable (free tier available)"),
    ("OpenAI", "GPT-3.5-turbo", "Set OPENAI_API_KEY environment variable"),
    ("Offline", "Rule-based system", "No API key needed — uses built-in medical knowledge base"),
]
for i, (mode, provider, setup) in enumerate(ai_modes):
    table.rows[i + 1].cells[0].text = mode
    table.rows[i + 1].cells[1].text = provider
    table.rows[i + 1].cells[2].text = setup

doc.add_paragraph()
doc.add_paragraph(
    "The system automatically detects available API keys and selects the appropriate provider. "
    "If no key is available, it falls back to the offline rule-based coach which provides "
    "structured recovery guidance based on the embedded medical knowledge base."
)

# ── Save ──
output_path = os.path.join(BASE_DIR, "Health_AI_Models_Documentation.docx")
try:
    doc.save(output_path)
except PermissionError:
    output_path = os.path.join(BASE_DIR, "Health_AI_Models_Documentation_new.docx")
    doc.save(output_path)
print(f"Document saved to: {output_path}")
